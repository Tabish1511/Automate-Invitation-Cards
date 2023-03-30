import os, docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

os.chdir('/home/tabish/Documents/LEARN_PYTHON/additional material for Automate the boring stuff/')

guestFile = open('guests.txt') # <<=== This text file should have guest names

os.chdir('/home/tabish/Documents/LEARN_PYTHON/PyProjects/')

gList = str(guestFile.read()).split('\n')

#doc = docx.Document('testProj.docx')

doc = docx.Document()


for i in range(len(gList)):
    para = doc.add_paragraph('It would be a pleasure to have the company of')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FreeMono'
    font.size = docx.shared.Pt(18)
    doc.paragraphs[(len(doc.paragraphs))-1].runs[0].italic = True
    #doc.paragraphs[0].runs[0].add_break()

    para = doc.add_paragraph(gList[i])
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style = doc.styles['Normal']
    font = style.font
    #font.name = 'FreeMono'
    font.size = docx.shared.Pt(18)
    doc.paragraphs[(len(doc.paragraphs))-1].runs[0].bold = True
    #doc.paragraphs[1].runs[0].add_break()
    
    para = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FreeMono'
    font.size = docx.shared.Pt(18)
    doc.paragraphs[(len(doc.paragraphs))-1].runs[0].italic = True
    #doc.paragraphs[2].runs[0].add_break()
    
    para = doc.add_paragraph('April 1st')
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style = doc.styles['Normal']
    font = style.font
    #font.name = 'FreeMono'
    font.size = docx.shared.Pt(18)
    #doc.paragraphs[(len(doc.paragraphs))-1].runs[0].italic = True
    #doc.paragraphs[3].runs[0].add_break()
    
    para = doc.add_paragraph("at 7 o'clock")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FreeMono'
    font.size = docx.shared.Pt(18)
    doc.paragraphs[(len(doc.paragraphs))-1].runs[0].italic = True
    doc.paragraphs[(len(doc.paragraphs))-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)




#os.chdir('/home/tabish/Documents/LEARN_PYTHON/PyProjects/')

doc.save('WordInvites.docx')
